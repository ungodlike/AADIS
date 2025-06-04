from crewai import Agent, Task, Crew, LLM
from crewai_tools import FileReadTool
import PyPDF2
from docx import Document
import json
from typing import Dict, List, Any

#text + table extraction for docx, only text extraction for pdf
#since only text is working for pdf tables so far based on my tests

class DocumentParsingTools:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    @staticmethod
    def extract_tables_from_docx(file_path: str) -> List[Dict]:
        """Extract tables from DOCX file"""
        try:
            doc = Document(file_path)
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append({
                    "table_id": i,
                    "data": table_data,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0
                })
            return tables
        except Exception as e:
            return [{"error": f"Error extracting tables: {str(e)}"}]

class DocumentProcessingCrew:
    def __init__(self):
        self.llm = LLM(model="groq/llama-3.3-70b-versatile")
        #init tools
        self.file_read_tool = FileReadTool()
        self.doc_tools = DocumentParsingTools()
        
        #text, table extraction agents
        self.text_extraction_agent = Agent(
            role="Text Extraction Specialist",
            goal="Extract and structure text content from documents with high accuracy",
            backstory="You are an expert in document analysis and text extraction. You specialize in identifying and extracting meaningful text content from various document formats while maintaining structure and context.",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )
        
        self.table_extraction_agent = Agent(
            role="Table Analysis Specialist", 
            goal="Identify, extract, and structure tabular data from documents",
            backstory="You are a specialist in identifying and extracting structured data from documents. You excel at finding tables, charts, and other structured information and converting them into usable formats.",
            verbose=True,
            allow_delegation=False,
            llm=self.llm
        )
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a document using specialized agents"""
        file_extension = filename.lower().split('.')[-1]
        
        #file based text extraction
        if file_extension == 'pdf':
            raw_text = self.doc_tools.extract_text_from_pdf(file_path)
            tables = []  #work on this!
        elif file_extension == 'docx':
            raw_text = self.doc_tools.extract_text_from_docx(file_path)
            tables = self.doc_tools.extract_tables_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        #text extract and table analysis tasks
        #still testing chunk size
        text_extraction_task = Task(
            description=f"""
            Analyze and structure the following text content from document '{filename}':
            
            {raw_text[:2000]}...  # Truncated for processing
            
            Your tasks:
            1. Clean and structure the text
            2. Identify key sections and headings
            3. Break text into meaningful chunks (max 250 words each) 
            4. Extract key information and metadata
            
            Return structured JSON with text_chunks, sections, and metadata.
            """,
            agent=self.text_extraction_agent,
            expected_output="JSON formatted text analysis with chunks and structure"
        )
        
        table_analysis_task = Task(
            description=f"""
            Analyze the tabular data extracted from document '{filename}':
            
            Tables found: {len(tables)}
            Table data: {json.dumps(tables, indent=2)}
            
            Your tasks:
            1. Analyze each table structure
            2. Identify headers and data types
            3. Look for captions if any
            3. Create meaningful descriptions for each table
            4. Structure data for easy querying
            
            Return structured JSON with table analysis and metadata.
            """,
            agent=self.table_extraction_agent,
            expected_output="JSON formatted table analysis with structure and descriptions"
        )
        
        crew = Crew(
            agents=[self.text_extraction_agent, self.table_extraction_agent],
            tasks=[text_extraction_task, table_analysis_task],
            verbose=True
        )
        
        result = crew.kickoff()
        
        #process end results
        try:
            #chunks for vector storage
            text_chunks = self._chunk_text(raw_text)
            
            return {
                "filename": filename,
                "text_chunks": text_chunks,
                "tables": tables,
                "raw_text": raw_text,
                "agent_analysis": str(result)
            }
        except Exception as e:
            #fallback
            return {
                "filename": filename,
                "text_chunks": self._chunk_text(raw_text),
                "tables": tables,
                "raw_text": raw_text,
                "agent_analysis": f"Agent processing failed: {str(e)}"
            }
    
    def _chunk_text(self, text: str, chunk_size: int = 500) -> List[str]:
        """Split text into chunks of specified size"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
